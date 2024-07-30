import pandas as pd
import numpy as np
import requests
from bs4 import BeautifulSoup
from transformers import RagTokenizer, RagSequenceForGeneration
from datasets import Dataset, load_from_disk
import faiss
import tempfile
import os
from docx import Document

# Fungsi untuk mengunduh konten dari URL
def download_content_from_url(url):
    try:
        response = requests.get(url)
        response.raise_for_status()  # Raise an error for bad responses
        soup = BeautifulSoup(response.text, 'html.parser')
        # Ambil semua teks dari elemen <p> sebagai contoh
        paragraphs = soup.find_all('p')
        text = ' '.join([para.get_text() for para in paragraphs])
        return text
    except Exception as e:
        print(f"Error downloading {url}: {e}")
        return ""

# Fungsi untuk memuat data dari file Excel, termasuk URL
def load_excel_data_from_urls(file_path, texts_file_path):
    df = pd.read_excel(file_path)
    # Mengunduh konten dari URL
    texts = []
    titles = df.get('Title', ['']*len(df))  # Menggunakan 'Title' jika ada
    for url in df['Url']:
        text = download_content_from_url(url)
        texts.append(text)
    
    # Simpan hasil ke file Excel
    output_df = pd.DataFrame({
        'Title': titles,
        'Text': texts
    })
    output_df.to_excel(texts_file_path, index=False)

    return Dataset.from_dict({
        'title': titles,
        'text': texts,
        'embeddings': [np.zeros(768).tolist() for _ in range(len(df))]  # Placeholder untuk embeddings
    })

# Fungsi untuk memuat data dari file Excel dengan teks berita
def load_excel_data_with_texts(file_path):
    df = pd.read_excel(file_path)
    return Dataset.from_dict({
        'title': df.get('Title', ['']*len(df)),
        'text': df['Text'].tolist(),
        'embeddings': [np.zeros(768).tolist() for _ in range(len(df))]  # Placeholder untuk embeddings
    })

# Simpan dataset ke disk
def save_dataset(dataset, path):
    dataset.save_to_disk(path)
    print(f"Dataset saved to {path}")

# Muat dataset dari disk
def load_dataset(path):
    dataset = load_from_disk(path)
    print(f"Dataset loaded from {path}")
    return dataset

# Inisialisasi tokenizer dan model RAG
rag_tokenizer = RagTokenizer.from_pretrained('facebook/rag-sequence-nq')
rag_model = RagSequenceForGeneration.from_pretrained('facebook/rag-sequence-nq')

# Path file Excel
urls_file_path = "E:/Program Skripsi Nlp/data/test.xlsx"  # File Excel dengan URL berita
texts_file_path = "E:/Program Skripsi Nlp/data/texts.xlsx"  # File Excel dengan teks berita
dataset_path = "E:/Program Skripsi Nlp/dataset"
index_path = "E:/Program Skripsi Nlp/faiss/index.faiss"

# Ambil teks dari URL dan simpan ke dataset
dataset_from_urls = load_excel_data_from_urls(urls_file_path, texts_file_path)
save_dataset(dataset_from_urls, dataset_path)

# Muat teks berita dari file Excel kedua dan perbarui dataset
dataset_from_texts = load_excel_data_with_texts(texts_file_path)
save_dataset(dataset_from_texts, dataset_path)

# Fungsi untuk membuat embedding dari teks menggunakan RAG
def embed_text(text):
    inputs = rag_tokenizer(text, return_tensors='pt', truncation=True)
    outputs = rag_model.question_encoder(**inputs)
    if isinstance(outputs, tuple):
        last_hidden_state = outputs[0]
    else:
        last_hidden_state = outputs.last_hidden_state
    return last_hidden_state.mean(dim=1).detach().numpy()  # Mengambil mean dari dimensi waktu

# Muat dataset dan simpan embeddings ke disk menggunakan FAISS
dataset = load_dataset(dataset_path)
embeddings = np.array([embed_text(item) for item in dataset['text']])
# Pastikan embeddings adalah 2D
if embeddings.ndim == 3:
    embeddings = embeddings.reshape(embeddings.shape[0], -1)

faiss_index = faiss.IndexFlatL2(embeddings.shape[1])  # Sesuaikan dengan dimensi embedding
faiss_index.add(embeddings.astype(np.float32))
faiss.write_index(faiss_index, index_path)

# Memuat kembali dataset dan index dari disk
dataset = load_from_disk(dataset_path)
faiss_index = faiss.read_index(index_path)

# Fungsi untuk membersihkan teks dari file DOCX
def clean_text_from_docx(file_path, unwanted_texts):
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    combined_text = "\n".join(full_text)
    for unwanted_text in unwanted_texts:
        combined_text = combined_text.replace(unwanted_text, "")
    return combined_text

# Fungsi untuk memproses file DOCX
def process_word_file(file):
    # Simpan file yang diunggah ke file sementara
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
        file.save(temp_file.name)
        temp_file_path = temp_file.name

    # Bersihkan teks dari file DOCX
    unwanted_texts = ["SCROLL TO CONTINUE WITH CONTENT"]
    cleaned_text = clean_text_from_docx(temp_file_path, unwanted_texts)
    
    # Buat query atau teks untuk diringkas
    query = cleaned_text
    
    # Pencarian Dokumen menggunakan RAG
    question_encodings = rag_tokenizer(query, return_tensors='pt')
    question_embeddings = rag_model.question_encoder(**question_encodings)
    if isinstance(question_embeddings, tuple):
        question_embeddings_np = question_embeddings[0].detach().numpy().astype(np.float32)
    else:
        question_embeddings_np = question_embeddings.last_hidden_state.detach().numpy().astype(np.float32)
    
    # Pencarian FAISS
    _, indices = faiss_index.search(question_embeddings_np, k=5)
    
    # Gabungkan teks relevan
    relevant_texts = [dataset['text'][i] for i in indices[0]]
    combined_text = " ".join(relevant_texts) + " " + query

    # Hasilkan ringkasan menggunakan model RAG
    inputs = rag_tokenizer(combined_text, return_tensors='pt', truncation=True)
    summary_ids = rag_model.generate(inputs['input_ids'], max_length=200, num_beams=4, early_stopping=True)
    summary_text = rag_tokenizer.decode(summary_ids[0], skip_special_tokens=True)

    # Buat file DOCX untuk ringkasan
    output_path = tempfile.mktemp(suffix='.docx')
    doc_out = Document()
    doc_out.add_paragraph(summary_text)
    doc_out.save(output_path)

    # Hapus file sementara
    os.remove(temp_file_path)

    return output_path